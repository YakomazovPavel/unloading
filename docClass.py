from docx import Document
from copy import deepcopy
from data_engine import getTemptureForOL, getPressureForOL, getFlowForOL, getLevelForOL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from data_engine import templates


def cell_style(cell, size):
    paragra_ph = cell.paragraphs[0]
    paragra_ph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = paragra_ph.runs[0].font
    font.name = 'Arial'
    font.size = Pt(size)


def unload(parametr, df_table_list_positions, df_ol_table):
    global savePath
    if parametr == 'Температура':
        document = Document(f'Шаблоны/ОЛ/Температура.docx')
        savePath = '193-РП-АТХ1.ОЛ1'
    elif parametr == 'Давление':
        document = Document(f'Шаблоны/ОЛ/Давление.docx')
        savePath = '193-РП-АТХ1.ОЛ2'
    elif parametr == 'Расход':
        document = Document(f'Шаблоны/ОЛ/Расход.docx')
        savePath = '193-РП-АТХ1.ОЛ3'
    elif parametr == 'Уровень':
        document = Document(f'Шаблоны/ОЛ/Уровень.docx')
        savePath = '193-РП-АТХ1.ОЛ4'
    else:
        return
    table_list_positions = document.tables[-2]

    for index, row in df_table_list_positions.iterrows():
        tb_row = table_list_positions.add_row()
        cells = tb_row.cells

        for _ in range(0, len(cells)):
            cell = cells[_]
            cell.text = str(row.iloc[_])
            cell_style(cell, 8)

    ol_table = document.tables[-1]

    for index, row in df_ol_table.iterrows():
        document.add_page_break()
        paragraph = document.add_paragraph()
        paragraph._p.addnext(deepcopy(ol_table._tbl))
        cur_table = document.tables[-1]
        for item in templates[parametr].keys():
            cell = cur_table.cell(*templates[parametr][item])
            cell.text = str(row[item])
            if templates[parametr][item][1] == 5:
                cell_style(cell, 10)
    ol_table._element.getparent().remove(ol_table._element)

    document.save(f'Выгрузка/Опросные листы/{savePath}.docx')


unload('Температура', *getTemptureForOL())
unload('Давление', *getPressureForOL())
unload('Расход', *getFlowForOL())
unload('Уровень', *getLevelForOL())
